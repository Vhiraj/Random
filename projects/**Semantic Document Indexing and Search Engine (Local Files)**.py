```python
import os
import re
import pickle
import logging
from typing import List, Dict, Tuple, Optional

# External Libraries - Install using:
# pip install sentence-transformers faiss-cpu PyPDF2 python-docx tqdm numpy
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from tqdm import tqdm

# --- Configuration ---
# You can change these parameters to suit your needs
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"  # A good balance of speed and performance
INDEX_FILE_PATH = "semantic_index.faiss"  # Path to save/load the FAISS index
METADATA_FILE_PATH = "semantic_metadata.pkl"  # Path to save/load document metadata
CHUNK_SIZE = 500  # Characters per text chunk
CHUNK_OVERLAP = 50  # Overlap between chunks to maintain context
TOP_K_RESULTS = 5  # Number of top results to return for a query

# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Helper Functions for Document Parsing ---

def _read_txt(filepath: str) -> str:
    """Reads text content from a .txt file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error reading .txt file {filepath}: {e}")
        return ""

def _read_pdf(filepath: str) -> str:
    """Reads text content from a .pdf file using PyPDF2."""
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        logger.error(f"Error reading .pdf file {filepath}: {e}")
        return ""

def _read_docx(filepath: str) -> str:
    """Reads text content from a .docx file using python-docx."""
    try:
        doc = DocxDocument(filepath)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error reading .docx file {filepath}: {e}")
        return ""

def _clean_text(text: str) -> str:
    """Cleans text by removing excessive whitespace and common noise."""
    text = re.sub(r'\s+', ' ', text).strip()  # Replace multiple spaces/newlines with single space
    # Add more cleaning steps if necessary (e.g., removing headers/footers, special characters)
    return text

def _chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """Splits text into overlapping chunks."""
    if not text:
        return []
    
    chunks = []
    current_pos = 0
    while current_pos < len(text):
        end_pos = min(current_pos + chunk_size, len(text))
        chunk = text[current_pos:end_pos]
        chunks.append(chunk)
        
        if end_pos == len(text):
            break
        
        current_pos += (chunk_size - chunk_overlap)
        # Ensure we don't go backwards or get stuck
        if current_pos < 0: 
            current_pos = 0
            
    return chunks

# --- Main Semantic Search Engine Class ---

class SemanticSearchEngine:
    """
    A local semantic search engine for indexing and searching text documents.

    It uses Sentence Transformers for generating embeddings and FAISS for efficient
    similarity search. Document content is chunked, and each chunk is indexed
    semantically. Supports .txt, .pdf, and .docx files.
    """

    def __init__(self, model_name: str = EMBEDDING_MODEL_NAME, 
                 index_path: str = INDEX_FILE_PATH, 
                 metadata_path: str = METADATA_FILE_PATH):
        """
        Initializes the SemanticSearchEngine.

        Args:
            model_name (str): The name of the Sentence Transformer model to use.
            index_path (str): File path to save/load the FAISS index.
            metadata_path (str): File path to save/load document metadata.
        """
        logger.info(f"Loading Sentence Transformer model: {model_name}...")
        self.model = SentenceTransformer(model_name)
        self.embedding_dim = self.model.get_sentence_embedding_dimension()
        logger.info(f"Model loaded. Embedding dimension: {self.embedding_dim}")

        self.index_path = index_path
        self.metadata_path = metadata_path
        self.faiss_index: Optional[faiss.IndexFlatL2] = None
        self.metadata: List[Dict] = []  # Stores [{'filepath': ..., 'chunk_text': ...}]

        self._load_index_and_metadata()

    def _load_index_and_metadata(self):
        """Loads the FAISS index and metadata from disk if they exist."""
        if os.path.exists(self.index_path) and os.path.exists(self.metadata_path):
            try:
                self.faiss_index = faiss.read_index(self.index_path)
                with open(self.metadata_path, 'rb') as f:
                    self.metadata = pickle.load(f)
                logger.info(f"Loaded existing FAISS index from {self.index_path} "
                            f"with {self.faiss_index.ntotal} vectors.")
                logger.info(f"Loaded metadata from {self.metadata_path} "
                            f"with {len(self.metadata)} entries.")
            except Exception as e:
                logger.error(f"Error loading index or metadata: {e}. Starting with empty index.")
                self._initialize_empty_index()
        else:
            logger.info("No existing index or metadata found. Initializing empty index.")
            self._initialize_empty_index()

    def _initialize_empty_index(self):
        """Initializes an empty FAISS index."""
        self.faiss_index = faiss.IndexFlatL2(self.embedding_dim)  # L2 distance for similarity
        self.metadata = []

    def _save_index_and_metadata(self):
        """Saves the current FAISS index and metadata to disk."""
        if self.faiss_index and self.faiss_index.ntotal > 0:
            try:
                faiss.write_index(self.faiss_index, self.index_path)
                with open(self.metadata_path, 'wb') as f:
                    pickle.dump(self.metadata, f)
                logger.info(f"FAISS index saved to {self.index_path}.")
                logger.info(f"Metadata saved to {self.metadata_path}.")
            except Exception as e:
                logger.error(f"Error saving index or metadata: {e}")
        else:
            logger.warning("No vectors in index to save.")

    def _process_file(self, filepath: str) -> Optional[List[Dict]]:
        """
        Reads, cleans, and chunks the content of a single file.

        Args:
            filepath (str): The path to the file.

        Returns:
            Optional[List[Dict]]: A list of dictionaries, where each dict contains
                                   'text' (chunk text) and 'filepath', or None if processing fails.
        """
        _, ext = os.path.splitext(filepath)
        ext = ext.lower()
        full_text = ""

        if ext == '.txt':
            full_text = _read_txt(filepath)
        elif ext == '.pdf':
            full_text = _read_pdf(filepath)
        elif ext == '.docx':
            full_text = _read_docx(filepath)
        else:
            logger.warning(f"Unsupported file type: {filepath}. Skipping.")
            return None

        if not full_text:
            logger.warning(f"No content extracted from {filepath}. Skipping.")
            return None

        cleaned_text = _clean_text(full_text)
        chunks = _chunk_text(cleaned_text, CHUNK_SIZE, CHUNK_OVERLAP)

        processed_chunks = []
        for chunk_text in chunks:
            processed_chunks.append({
                'filepath': filepath,
                'chunk_text': chunk_text
            })
        return processed_chunks

    def index_documents(self, directory_path: str):
        """
        Indexes all supported documents in a given directory.

        Args:
            directory_path (str): The path to the directory containing documents.
        """
        if not os.path.isdir(directory_path):
            logger.error(f"Directory not found: {directory_path}")
            return

        logger.info(f"Starting document indexing in directory: {directory_path}")
        
        # Collect all eligible files first
        all_files = []
        for root, _, files in os.walk(directory_path):
            for file in files:
                filepath = os.path.join(root, file)
                if any(filepath.lower().endswith(ext) for ext in ['.txt', '.pdf', '.docx']):
                    all_files.append(filepath)

        if not all_files:
            logger.warning(f"No supported documents found in {directory_path}. "
                           "Supported types: .txt, .pdf, .docx")
            return

        new_embeddings = []
        new_metadata_entries = []

        for filepath in tqdm(all_files, desc="Processing documents"):
            processed_chunks = self._process_file(filepath)
            if processed_chunks:
                # Check if document already indexed (simple check by filepath, could be more robust)
                is_indexed = any(m['filepath'] == filepath for m in self.metadata)
                if is_indexed:
                    logger.info(f"Document {filepath} already indexed. Skipping.")
                    continue

                chunk_texts = [chunk['chunk_text'] for chunk in processed_chunks]
                
                try:
                    # Generate embeddings for all chunks of the current document
                    embeddings = self.model.encode(chunk_texts, convert_to_numpy=True)
                    new_embeddings.extend(embeddings)
                    new_metadata_entries.extend(processed_chunks)
                except Exception as e:
                    logger.error(f"Error generating embeddings for {filepath}: {e}")
        
        if new_embeddings:
            logger.info(f"Adding {len(new_embeddings)} new embeddings to FAISS index.")
            self.faiss_index.add(np.array(new_embeddings).astype('float32'))
            self.metadata.extend(new_metadata_entries)
            self._save_index_and_metadata()
            logger.info(f"Indexing complete. Total vectors in index: {self.faiss_index.ntotal}")
        else:
            logger.info("No new documents processed or no new embeddings generated.")

    def search(self, query: str, top_k: int = TOP_K_RESULTS) -> List[Tuple[float, str, str]]:
        """
        Searches the indexed documents for semantically similar content.

        Args:
            query (str): The search query string.
            top_k (int): The number of top results to return.

        Returns:
            List[Tuple[float, str, str]]: A list of tuples, where each tuple contains
                                         (similarity_score, document_filepath, relevant_text_chunk).
        """
        if self.faiss_index is None or self.faiss_index.ntotal == 0:
            logger.warning("Index is empty. Please index documents first.")
            return []

        try:
            query_embedding = self.model.encode([query], convert_to_numpy=True).astype('float32')
            
            # D, I = distances, indices
            distances, indices = self.faiss_index.search(query_embedding, top_k)
            
            results = []
            for i in range(top_k):
                # FAISS uses L2 distance, so smaller distance means higher similarity.
                # To convert to a similarity score (0 to 1), we can use a simple inverse or exponential.
                # For now, let's just return the inverse of the distance (or negative distance)
                # to show that smaller value means closer (better).
                # A common approach for semantic similarity is to use cosine similarity,
                # but FAISS IndexFlatL2 uses Euclidean distance.
                # If we used IndexFlatIP, it would be dot product which directly relates to cosine similarity.
                # For L2, a simple way to get a 'score' is `1 / (1 + distance)`.
                
                faiss_idx = indices[0][i]
                if faiss_idx == -1: # FAISS returns -1 for empty slots if k is too high
                    continue

                distance = distances[0][i]
                
                # Check bounds for metadata
                if 0 <= faiss_idx < len(self.metadata):
                    metadata_entry = self.metadata[faiss_idx]
                    filepath = metadata_entry['filepath']
                    chunk_text = metadata_entry['chunk_text']
                    
                    # Convert L2 distance to a more intuitive similarity score (higher is better)
                    # This is an approximation; true cosine similarity requires IndexFlatIP or normalization.
                    # For L2, smaller distance means more similar. We can express this as:
                    # similarity = 1 - (distance / max_possible_distance)
                    # Or a simpler inverse:
                    similarity_score = 1.0 / (1.0 + distance) # Higher score for smaller distance

                    results.append((similarity_score, filepath, chunk_text))
                else:
                    logger.warning(f"FAISS index {faiss_idx} out of bounds for metadata.")
            
            # Sort by similarity score (descending)
            results.sort(key=lambda x: x[0], reverse=True)
            return results

        except Exception as e:
            logger.error(f"Error during search: {e}")
            return []
            
    def clear_index(self):
        """Clears the FAISS index and metadata, and removes saved files."""
        logger.info("Clearing existing index and metadata.")
        self._initialize_empty_index() # Reset the in-memory index
        
        # Remove saved files
        if os.path.exists(self.index_path):
            os.remove(self.index_path)
            logger.info(f"Removed {self.index_path}")
        if os.path.exists(self.metadata_path):
            os.remove(self.metadata_path)
            logger.info(f"Removed {self.metadata_path}")
        
        logger.info("Index and metadata cleared successfully.")


# --- Main Execution Block ---

if __name__ == "__main__":
    # --- 1. Setup Sample Documents (for demonstration) ---
    sample_docs_dir = "sample_documents"
    os.makedirs(sample_docs_dir, exist_ok=True)

    # Create a .txt file
    with open(os.path.join(sample_docs_dir, "report_on_ai.txt"), "w", encoding='utf-8') as f:
        f.write("Artificial intelligence (AI) is rapidly transforming various industries. "
                "Machine learning, a subset of AI, enables systems to learn from data without explicit programming. "
                "Deep learning, a more advanced form, uses neural networks with many layers. "
                "Natural Language Processing (NLP) is a key AI area focused on human-computer language interaction. "
                "Robotics often integrates AI for autonomous tasks. The future of AI promises significant advancements.")

    # Create another .txt file
    with open(os.path.join(sample_docs_dir, "future_of_work.txt"), "w", encoding='utf-8') as f:
        f.write("The future of work is being shaped by automation and remote collaboration. "
                "Gig economy platforms are becoming more prevalent. "
                "Employees need to adapt to new skills and lifelong learning. "
                "Flexible work arrangements are increasing, fostering a better work-life balance. "
                "Technological advancements are driving these changes.")
                
    # Create a dummy .pdf file (PyPDF2 cannot create, so we'll make a simple one)
    # For a real PDF, you'd generate it with a library like ReportLab or manually place one.
    # Here, we'll just acknowledge if a real PDF existed, it would be parsed.
    # For a full test, manually place a PDF.
    # This example will assume an empty PDF or one that needs manual placement for robust test.
    # Example for creating a dummy PDF (requires reportlab)
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        pdf_path = os.path.join(sample_docs_dir, "project_summary.pdf")
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "Project Gamma: Summary Report.")
        c.drawString(100, 730, "This project focuses on sustainable energy solutions.")
        c.drawString(100, 710, "It involves solar power and wind turbine technologies.")
        c.drawString(100, 690, "The goal is to reduce carbon emissions and promote green energy.")
        c.save()
        logger.info(f"Created dummy PDF: {pdf_path}")
    except ImportError:
        logger.warning("ReportLab not installed. Cannot create dummy PDF. "
                       "Please place a 'project_summary.pdf' manually for full PDF testing.")
        with open(os.path.join(sample_docs_dir, "project_summary.pdf.txt"), "w", encoding='utf-8') as f:
             f.write("This is a placeholder for a PDF. Imagine it contains information about project summaries, sustainable energy, and green initiatives.")


    # Create a dummy .docx file (requires python-docx)
    try:
        from docx import Document as DocxDocument
        docx_path = os.path.join(sample_docs_dir, "meeting_minutes.docx")
        doc = DocxDocument()
        doc.add_heading('Meeting Minutes - Q3 Planning', level=1)
        doc.add_paragraph('Date: October 26, 2023')
        doc.add_paragraph('Attendees: John Doe, Jane Smith, Alice Johnson')
        doc.add_paragraph('Key discussions included budget allocation for new AI projects '
                          'and strategies for expanding market reach in automation.')
        doc.add_paragraph('Action Item: Jane to follow up on AI integration proposals.')
        doc.save(docx_path)
        logger.info(f"Created dummy DOCX: {docx_path}")
    except ImportError:
        logger.warning("python-docx not installed. Cannot create dummy DOCX. "
                       "Please place a 'meeting_minutes.docx' manually for full DOCX testing.")
        with open(os.path.join(sample_docs_dir, "meeting_minutes.docx.txt"), "w", encoding='utf-8') as f:
            f.write("This is a placeholder for a DOCX file. It might contain meeting minutes about AI projects and automation strategies.")

    # --- 2. Initialize and Use the Search Engine ---
    print("\n--- Initializing Semantic Search Engine ---")
    engine = SemanticSearchEngine()

    # --- 3. Index Documents ---
    # You can call this multiple times; it will only index new documents.
    # To re-index everything, call engine.clear_index() first.
    print(f"\n--- Indexing documents in '{sample_docs_dir}' ---")
    engine.index_documents(sample_docs_dir)

    # --- 4. Perform Searches ---
    print("\n--- Performing Searches ---")
    while True:
        query = input("\nEnter your search query (or 'exit' to quit): ").strip()
        if query.lower() == 'exit':
            break

        if not query:
            print("Query cannot be empty.")
            continue

        print(f"\nSearching for: '{query}'")
        results = engine.search(query, top_k=TOP_K_RESULTS)

        if results:
            print(f"Found {len(results)} relevant results:")
            for i, (score, filepath, chunk_text) in enumerate(results):
                print(f"\n--- Result {i+1} (Score: {score:.4f}) ---")
                print(f"File: {os.path.basename(filepath)}")
                print(f"Relevant Text Snippet: \"{chunk_text.strip()}\"")
        else:
            print("No results found for your query.")

    print("\n--- Application Finished ---")

    # Optional: Clean up sample documents
    # import shutil
    # if os.path.exists(sample_docs_dir):
    #     shutil.rmtree(sample_docs_dir)
    #     print(f"Cleaned up '{sample_docs_dir}' directory.")
    # # Optional: Clean up index files
    # if os.path.exists(INDEX_FILE_PATH):
    #     os.remove(INDEX_FILE_PATH)
    # if os.path.exists(METADATA_FILE_PATH):
    #     os.remove(METADATA_FILE_PATH)
    #     print("Cleaned up FAISS index and metadata files.")
```